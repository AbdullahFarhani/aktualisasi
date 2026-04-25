from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_notula_docx():
    doc = Document()

    # Header / KOP (Optional but makes it look professional)
    title = doc.add_heading('NOTULA HASIL KONSULTASI DENGAN MENTOR', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('KEGIATAN AKTUALISASI CPNS KODAM V/BRAWIJAYA', 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Table for Metadata
    table_meta = doc.add_table(rows=4, cols=2)
    table_meta.columns[0].width = Inches(1.5)
    
    cells = table_meta.rows[0].cells
    cells[0].text = "Hari / Tanggal"
    cells[1].text = ": Senin, 20 April 2026"
    
    cells = table_meta.rows[1].cells
    cells[0].text = "Waktu"
    cells[1].text = ": 09.00 - 10.30 WIB"
    
    cells = table_meta.rows[2].cells
    cells[0].text = "Tempat"
    cells[1].text = ": Ruang Kerja Infolahtadam V/Brawijaya"
    
    cells = table_meta.rows[3].cells
    cells[0].text = "Topik"
    cells[1].text = ": Optimalisasi Sistem Intelijen Berita dan Penanganan Kendala Jaringan (Google News Block)"

    doc.add_paragraph("\n")

    # 1. Poin-Poin Masukan/Arahan Mentor
    doc.add_heading('A. MASUKAN DAN ARAHAN MENTOR', level=1)
    
    inputs = [
        "Mentor mengapresiasi kemajuan pengembangan sistem, khususnya pada akurasi identifikasi berita ancaman menggunakan model AI Groq dan Gemini.",
        "Mentor memberikan masukan agar sistem harus lebih tangguh dalam menghadapi pemblokiran akses (Error 429) dari pihak penyedia data (Google News).",
        "Pentingnya fitur untuk mendapatkan tautan asli (Direct Link) agar unit intelijen dapat memverifikasi sumber berita secara langsung tanpa hambatan redirect.",
        "Mentor menyarankan penambahan fitur rotasi identitas browser (User-Agent) dan penggunaan proxy jika diperlukan untuk menjaga stabilitas crawling.",
        "Segala bentuk API Key dan konfigurasi jaringan harus disimpan dengan aman dan terstruktur."
    ]
    
    for item in inputs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    # 2. Rencana Tindak Lanjut
    doc.add_heading('B. RENCANA TINDAK LANJUT (KESPAKATAN)', level=1)
    
    follow_ups = [
        "Melakukan implementasi fitur 'Auto-Proxy Harvester' yang mampu mencari jalur aman secara otomatis saat terdeteksi pemblokiran.",
        "Mengembangkan mekanisme 'Multi-RPC Fallback' untuk memecahkan enkripsi tautan Google News v2026.",
        "Melakukan uji coba (stress test) pada portal berita yang memiliki proteksi ketat seperti Surya.co.id, Detik.com, dan Tribunnews.",
        "Menyelesaikan Dokumen System Requirements (SRD) sebagai dasar teknis sistem.",
        "Melaporkan hasil uji coba stabilitas sistem pada sesi konsultasi minggu depan."
    ]
    
    for item in follow_ups:
        p = doc.add_paragraph(style='List Number')
        p.add_run(item)

    # Signatures
    doc.add_paragraph("\n\n")
    table_sign = doc.add_table(rows=1, cols=2)
    table_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    cells = table_sign.rows[0].cells
    
    # Left Side: Participant
    p_part = cells[0].add_paragraph()
    p_part.add_run("Peserta Aktualisasi,\n\n\n\n( Abdullah Farhani )\nNIP.").bold = True
    p_part.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Right Side: Mentor
    p_ment = cells[1].add_paragraph()
    p_ment.add_run("Mentor,\n\n\n\n( ________________________ )\nNIP.").bold = True
    p_ment.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    file_name = 'Notula_Konsultasi_Mentor_Abdullah_Farhani.docx'
    doc.save(file_name)
    print(f"File notula berhasil dibuat: {file_name}")

if __name__ == "__main__":
    create_notula_docx()
