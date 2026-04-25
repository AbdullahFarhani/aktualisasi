from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_srd_docx():
    doc = Document()

    # Title
    title = doc.add_heading('DOKUMEN SPESIFIKASI KEBUTUHAN SISTEM (SYSTEM REQUIREMENTS DOCUMENT)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('PROYEK: SISTEM INTELIJEN BERITA DIGITAL (NEWS INTELLIGENCE SYSTEM) v5.45', 1)
    
    p = doc.add_paragraph()
    p.add_run('Unit Kerja: ').bold = True
    p.add_run('Kodam V/Brawijaya\n')
    p.add_run('Penyusun: ').bold = True
    p.add_run('Abdullah Farhani\n')
    p.add_run('Tanggal: ').bold = True
    p.add_run('20 April 2026')

    doc.add_heading('1. PENDAHULUAN', level=1)
    doc.add_heading('1.1 Tujuan Dokumen', level=2)
    doc.add_paragraph(
        "Dokumen ini disusun untuk mendefinisikan seluruh kebutuhan teknis dan operasional dari Sistem Intelijen Berita Digital (v5.45). "
        "Dokumen ini berfungsi sebagai bukti fisik (eviden) dalam Laporan Aktualisasi untuk menunjukkan proses perencanaan dan analisis sistem yang sistematis."
    )

    doc.add_heading('1.2 Lingkup Sistem', level=2)
    doc.add_paragraph(
        "Sistem ini dirancang untuk melakukan pemantauan berita secara otomatis dari berbagai portal portal berita nasional dan regional Jawa Timur "
        "melalui integrasi Google News RSS, pemrosesan bahasa alami (NLP) menggunakan AI, dan sistem peringatan dini (Early Warning System) melalui Telegram."
    )

    doc.add_heading('2. KEBUTUHAN FUNGSIONAL (FUNCTIONAL REQUIREMENTS)', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Deskripsi Kebutuhan'
    hdr_cells[2].text = 'Detail Teknis'

    data_fr = [
        ('FR-01', 'Automated News Harvesting', 'Sistem harus dapat mengambil data berita secara berkala (crawling) berdasarkan kata kunci ancaman (threat keywords) dan wilayah geografis (38 Kab/Kota di Jawa Timur).'),
        ('FR-02', 'Google News URL Decoding', 'Sistem harus mampu memecahkan (decrypt) tautan terenkripsi Google News (pola AU_yqL / CBM) untuk mendapatkan URL asli dari portal berita penerbit.'),
        ('FR-03', 'AI Threat Profiling', 'Sistem harus mampu menganalisis konten berita menggunakan model AI (Groq/Gemini/IndoRoBERTa) untuk menentukan kategori ancaman (Positif, Netral, atau Ancaman).'),
        ('FR-04', 'Editorial Metadata Extraction', 'Sistem harus mampu mengekstraksi data redaksi (Reporter, Editor, Penanggung Jawab) dan alamat kantor berita secara otomatis dari halaman web.'),
        ('FR-05', 'Early Warning Alerting', 'Sistem harus mengirimkan notifikasi instan ke grup Telegram Intelijen jika ditemukan berita berkategori "Ancaman" (⚠️).'),
        ('FR-06', 'Google News Fallback', 'Jika teks asli artikel tidak dapat dijangkau, sistem harus mampu melakukan ekstraksi deskripsi pintar dari cache Google News sebagai jalur cadangan.')
    ]

    for id_req, desc, tech in data_fr:
        row_cells = table.add_row().cells
        row_cells[0].text = id_req
        row_cells[1].text = desc
        row_cells[2].text = tech

    doc.add_heading('3. KEBUTUHAN NON-FUNGSIONAL (NON-FUNCTIONAL REQUIREMENTS)', level=1)
    table_nfr = doc.add_table(rows=1, cols=3)
    table_nfr.style = 'Table Grid'
    hdr_nfr = table_nfr.rows[0].cells
    hdr_nfr[0].text = 'ID'
    hdr_nfr[1].text = 'Kategori'
    hdr_nfr[2].text = 'Spesifikasi Kebutuhan'

    data_nfr = [
        ('NFR-01', 'Reliability (Resilience)', 'Sistem harus memiliki mekanisme Lapis 1-5 untuk pemulihan tautan, termasuk penggunaan mesin pencari cadangan (DuckDuckGo Lite/Yandex) jika Google melakukan pemblokiran.'),
        ('NFR-02', 'Availability (Proxy)', 'Sistem harus mampu melakukan perburuan proxy otomatis (Auto-Proxy Harvester) dan rotasi identitas browser untuk menghindari limitasi akses (Error 429).'),
        ('NFR-03', 'Performance', 'Waktu rata-rata pemrosesan satu artikel (crawling hingga AI profiling) tidak boleh lebih dari 15 detik pada koneksi internet standar.'),
        ('NFR-04', 'Security', 'Penggunaan DNS-over-HTTPS (DoH) via Cloudflare untuk menghindari pemblokiran DNS ISP (Internet Positif) dan enkripsi data API Key.')
    ]

    for id_req, cat, spec in data_nfr:
        row_cells = table_nfr.add_row().cells
        row_cells[0].text = id_req
        row_cells[1].text = cat
        row_cells[2].text = spec

    doc.add_heading('4. KEBUTUHAN ANTARMUKA (INTERFACE REQUIREMENTS)', level=1)
    doc.add_paragraph("1. Antarmuka Eksternal:", style='List Bullet')
    doc.add_paragraph("Google News RSS API: Sebagai sumber data utama.", style='List Bullet 2')
    doc.add_paragraph("Telegram Bot API: Sebagai media pengiriman laporan intelijen.", style='List Bullet 2')
    doc.add_paragraph("Groq/Gemini API: Sebagai mesin pemrosesan kecerdasan buatan.", style='List Bullet 2')
    
    doc.add_paragraph("2. Antarmuka Pengguna:", style='List Bullet')
    doc.add_paragraph("Sistem berbasis Command Line Interface (CLI) untuk efisiensi server.", style='List Bullet 2')
    doc.add_paragraph("Dashboard berbasis Chat (Telegram) untuk pengguna akhir (User).", style='List Bullet 2')

    doc.add_heading('5. KEBUTUHAN PERANGKAT (HARDWARE & SOFTWARE REQUIREMENTS)', level=1)
    doc.add_paragraph("Sistem Operasi: Linux (Ubuntu/Debian recommended).", style='List Bullet')
    doc.add_paragraph("Bahasa Pemrograman: Python 3.10+.", style='List Bullet')
    doc.add_paragraph("Library Utama: requests, BeautifulSoup4, Newspaper3k, Selectolax, Groq-SDK.", style='List Bullet')
    doc.add_paragraph("Spesifikasi Minimum: 1 vCPU, 2GB RAM, 10GB Disk Space.", style='List Bullet')

    doc.add_paragraph("\n\n")
    p_sign = doc.add_paragraph()
    p_sign.add_run("Mengetahui,\nMentor Latsar CPNS\n\n\n\n(________________________)\nNIP.").bold = True
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save('System_Requirements_Abdullah_Farhani.docx')
    print("File docx berhasil dibuat: System_Requirements_Abdullah_Farhani.docx")

if __name__ == "__main__":
    create_srd_docx()
