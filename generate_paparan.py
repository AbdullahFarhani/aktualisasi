from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_paparan_docx():
    doc = Document()

    # Title Slide Equivalent
    title = doc.add_heading('KONSEP PAPARAN RENCANA PENGEMBANGAN SISTEM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('OPTIMALISASI MONITORING MEDIA DIGITAL MELALUI NEWS INTELLIGENCE SYSTEM v5.45', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("KODAM V/BRAWIJAYA").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("_" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # I. Latar Belakang & Isu
    doc.add_heading('I. LATAR BELAKANG DAN ISU STRATEGIS', level=1)
    p1 = doc.add_paragraph()
    p1.add_run("Berdasarkan pengamatan di lapangan, ditemukan beberapa isu utama dalam proses monitoring media di lingkungan unit kerja:").bold = True
    
    issues = [
        "Proses monitoring berita masih dilakukan secara konvensional/manual, sehingga membutuhkan waktu dan personel yang banyak.",
        "Potensi keterlambatan dalam mendeteksi berita negatif atau ancaman (early warning) yang dapat berdampak pada kondusivitas wilayah.",
        "Adanya kendala teknis berupa pemblokiran akses (Rate Limit/429) dari portal berita besar yang menghambat perolehan data intelijen secara cepat.",
        "Belum tersedianya sistem klasifikasi berita berbasis kecerdasan buatan (AI) yang mampu memilah antara berita positif, netral, dan ancaman secara otomatis."
    ]
    for issue in issues:
        doc.add_paragraph(issue, style='List Bullet')

    # II. Gagasan Pemecahan Isu
    doc.add_heading('II. GAGASAN PEMECAHAN ISU', level=1)
    p2 = doc.add_paragraph()
    p2.add_run("Solusi yang diajukan adalah pengembangan Sistem Intelijen Berita Digital (News Intelligence System) v5.45 dengan fitur-fitur unggulan:").bold = True
    
    solutions = [
        "Automated Crawling: Pemantauan otomatis 38 Kabupaten/Kota di Jawa Timur selama 24/7.",
        "AI-Powered Profiling: Menggunakan model Llama-3/Gemini untuk analisis sentimen dan klasifikasi ancaman secara real-time.",
        "Genius Network Resilience: Implementasi Auto-Proxy Harvester dan Multi-RPC Decoder untuk menembus proteksi portal berita.",
        "Early Warning System: Notifikasi instan ke Telegram Intelijen untuk respon cepat terhadap isu menonjol."
    ]
    for sol in solutions:
        doc.add_paragraph(sol, style='List Bullet')

    # III. Tahapan Kegiatan Aktualisasi
    doc.add_heading('III. TAHAPAN KEGIATAN AKTUALISASI', level=1)
    stages = [
        ("Tahap 1: Perencanaan & Analisis", "Melakukan konsultasi dengan mentor dan menyusun System Requirements Document (SRD)."),
        ("Tahap 2: Pengembangan Inti", "Membangun mesin crawler dan mengintegrasikan API Kecerdasan Buatan (Groq/Gemini)."),
        ("Tahap 3: Stabilisasi & Ketangguhan", "Implementasi fitur bypass pemblokiran dan pemulihan tautan asli (Resilience System)."),
        ("Tahap 4: Uji Coba & Evaluasi", "Running sistem secara penuh dan melakukan perbaikan berdasarkan hasil monitoring harian."),
        ("Tahap 5: Pelaporan & Sosialisasi", "Penyusunan laporan akhir dan penyerahan manual penggunaan sistem.")
    ]
    
    for stage, desc in stages:
        p = doc.add_paragraph()
        p.add_run(f"{stage}: ").bold = True
        p.add_run(desc)

    # IV. Target Capaian Aktualisasi
    doc.add_heading('IV. TARGET CAPAIAN DAN MANFAAT', level=1)
    targets = [
        "Efisiensi Waktu: Mengurangi waktu monitoring manual hingga 80%.",
        "Kecepatan Respon: Informasi ancaman sampai ke pimpinan dalam waktu < 1 menit setelah terbit di media.",
        "Akurasi Data: Klasifikasi berita oleh AI mencapai tingkat akurasi > 90%.",
        "Ketangguhan Akses: Sistem mampu menembus blokir jaringan pada 10+ portal berita utama nasional."
    ]
    for target in targets:
        doc.add_paragraph(target, style='List Bullet')

    # Footer
    doc.add_paragraph("\n")
    p_end = doc.add_paragraph()
    p_end.add_run("Demikian konsep paparan ini disusun untuk mendapatkan arahan dan persetujuan lebih lanjut.\n\nSurabaya, 20 April 2026\n\nPenulis,\n\n\n\nAbdullah Farhani")
    p_end.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Save
    file_name = 'Konsep_Paparan_Rencana_Sistem_Abdullah_Farhani.docx'
    doc.save(file_name)
    print(f"File paparan berhasil dibuat: {file_name}")

if __name__ == "__main__":
    create_paparan_docx()
